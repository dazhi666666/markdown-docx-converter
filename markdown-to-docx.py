from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from markdown_it import MarkdownIt


def process_markdown_runs(paragraph, text):
    """通用函数：解析文本中的**加粗**和*斜体*，并添加到段落"""
    # 分割加粗语法**
    parts = text.split("**")
    is_bold = False
    for part in parts:
        # 进一步分割斜体语法*
        sub_parts = part.split("*")
        for i, sub_part in enumerate(sub_parts):
            if not sub_part:
                continue
            run = paragraph.add_run(sub_part)
            run.bold = is_bold
            # 斜体判断（奇数索引为斜体部分）
            run.italic = i % 2 == 1
        # 切换加粗状态
        is_bold = not is_bold


def markdown_to_docx(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()
    md = MarkdownIt()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            new_doc.add_paragraph()
            continue

        # 处理标题（如 #### 标题）
        if text.startswith("#"):
            # 提取标题级别和内容
            header_part = text.split(maxsplit=1)
            level = len(header_part[0].lstrip('#'))  # 计算#数量
            level = max(1, min(6, level))  # 限制1-6级标题
            header_text = header_part[1] if len(header_part) > 1 else ""

            # 添加标题段落并处理格式
            heading_para = new_doc.add_heading(level=level)
            process_markdown_runs(heading_para, header_text)

        # 处理列表项
        elif text.startswith(("* ", "- ")):
            list_para = new_doc.add_paragraph(style="List Bullet")
            process_markdown_runs(list_para, text.lstrip("* -"))

        # 普通段落
        else:
            new_para = new_doc.add_paragraph()
            process_markdown_runs(new_para, text)

    new_doc.save(output_path)


# 调用示例
markdown_to_docx("input.docx", "output_optimized.docx")